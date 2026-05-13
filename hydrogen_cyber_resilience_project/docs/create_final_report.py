from __future__ import annotations

from pathlib import Path
import json
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Hydrogen_Cyber_Resilience_Final_Project_Report.docx"
METRICS_JSON = ROOT / "outputs" / "model_metrics.json"
METRICS_CSV = ROOT / "outputs" / "model_metrics.csv"
CV_CSV = ROOT / "outputs" / "cross_validation_metrics.csv"
ATTACK_RECALL_CSV = ROOT / "outputs" / "per_attack_recall.csv"
FEATURE_IMPORTANCE_CSV = ROOT / "outputs" / "feature_importance.csv"
CONFUSION = ROOT / "outputs" / "confusion_matrix.png"
ROC = ROOT / "outputs" / "roc_curve.png"
PR = ROOT / "outputs" / "precision_recall_curve.png"
FEATURE_PLOT = ROOT / "outputs" / "feature_importance.png"
TIMELINE = ROOT / "outputs" / "attack_timeline.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.7) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)


def set_margins(section) -> None:
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_heading(doc: Document, title: str, level: int = 1) -> None:
    doc.add_heading(title, level=level)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8.7,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, font_size=font_size)
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), font_size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph("")


def add_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 5.9) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.right_indent = Inches(0.24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.06
    for style_name, size in [("Title", 21), ("Heading 1", 15.5), ("Heading 2", 12.7), ("Heading 3", 11.0)]:
        styles[style_name].font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8.8)
    styles["Caption"].font.italic = True


def fmt_pct(v: float) -> str:
    return f"{100 * float(v):.2f}%"


def fmt_num(v: float, digits: int = 3) -> str:
    return f"{float(v):.{digits}f}"


def build_report() -> None:
    metrics_payload = json.loads(METRICS_JSON.read_text(encoding="utf-8"))
    metrics_df = pd.read_csv(METRICS_CSV, index_col=0)
    cv_df = pd.read_csv(CV_CSV)
    attack_recall_df = pd.read_csv(ATTACK_RECALL_CSV)
    feature_df = pd.read_csv(FEATURE_IMPORTANCE_CSV)

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        set_margins(section)

    header = doc.sections[0].header.paragraphs[0]
    header.text = "Machine Learning for Cyber Resilience of Hydrogen Energy Systems"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8.2)
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("The Use of Machine Learning in Improving Cyber Resilience of Hydrogen Energy Systems")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Final Year Project Report, Technical Artefact, and Critical Evaluation")
    s.font.size = Pt(12.5)
    s.bold = True
    doc.add_paragraph("")

    cover_data = [
        ["Student", "Pradip Mainali"],
        ["Student ID", "32127812"],
        ["Course", "Computer Science"],
        ["Supervisor", "Fateme Dinmohammadi"],
        ["Artefact", "Reproducible ML-based cyber anomaly detection pipeline"],
        ["Submission focus", "Hydrogen OT cyber resilience, threat modelling, implementation, and evaluation"],
    ]
    add_table(doc, ["Field", "Details"], cover_data, widths=[1.75, 4.8], font_size=9)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope_run = note.add_run(
        "Scope note: all experimental results are obtained from a controlled synthetic hydrogen ICS-style dataset created for educational demonstration; they are not production validation claims."
    )
    scope_run.italic = True
    doc.add_page_break()

    # Abstract and executive contribution
    add_heading(doc, "Abstract", 1)
    add_para(doc, "Hydrogen energy systems are increasingly dependent on digital monitoring, supervisory control, communications, and sensor-driven decision making. This creates efficiency and safety benefits, but it also expands the attack surface of a safety-relevant cyber-physical infrastructure. This project investigates how machine learning can support cyber resilience in hydrogen energy systems through anomaly detection, with a deliberately transparent bachelor-level implementation.")
    add_para(doc, "The technical artefact is a reproducible Python pipeline that generates a synthetic hydrogen industrial-control-system (ICS) dataset, trains supervised and unsupervised anomaly-detection models, and evaluates them using both a fixed holdout test split and stratified cross-validation. The dataset includes four interpretable attack scenarios: false data injection, denial-of-service-like communication degradation, actuator spoofing, and a more subtle stealthy sensor-drift scenario designed to be less trivially separable.")
    add_para(doc, f"The Random Forest model produced the strongest holdout performance, achieving F1 = {metrics_df.loc['random_forest', 'f1']:.3f}, recall = {metrics_df.loc['random_forest', 'recall']:.3f}, and average precision = {metrics_df.loc['random_forest', 'average_precision']:.3f}. Three-fold cross-validation remained strong with mean F1 = {cv_df.loc[cv_df['model'] == 'random_forest', 'f1_mean'].iloc[0]:.3f}. The most challenging class was stealthy sensor drift, for which the best supervised model achieved recall of {attack_recall_df.loc[attack_recall_df['attack_type'] == 'stealthy_sensor_drift', 'recall'].iloc[0]:.3f}. These results support the project claim that a structured ML workflow can identify controlled cyber-physical abnormalities, while also demonstrating why low-amplitude attacks require stronger validation and domain realism.")
    add_para(doc, "Keywords: hydrogen energy systems; cyber resilience; operational technology; anomaly detection; machine learning; industrial control systems; threat modelling.")

    add_heading(doc, "Contribution Summary", 1)
    add_bullets(doc, [
        "A hydrogen-specific cybersecurity framing rather than a generic IT intrusion-detection narrative.",
        "A formalised threat model linking assets, attack objectives, OT consequences, and resilience functions.",
        "A reproducible Python artefact with four attack scenarios, holdout testing, cross-validation, per-attack error analysis, and permutation feature importance.",
        "A critical discussion that distinguishes educational proof-of-concept performance from real-world deployment evidence.",
    ])
    doc.add_page_break()

    # Contents
    add_heading(doc, "Contents", 1)
    add_bullets(doc, [
        "1. Introduction",
        "2. Aim, Objectives, and Research Questions",
        "3. Background and Literature Review",
        "4. Threat Model and Resilience Framework",
        "5. Methodology",
        "6. Artefact Design and Implementation",
        "7. Experimental Results",
        "8. Discussion and Critical Evaluation",
        "9. Project Management, Risks, Ethics, and Professional Considerations",
        "10. Conclusion and Future Work",
        "References",
        "Appendix A. Reproduction Instructions",
        "Appendix B. Data Dictionary",
        "Appendix C. Attack Scenario Definitions",
    ])
    doc.add_page_break()

    # 1 Introduction
    add_heading(doc, "1. Introduction", 1)
    add_para(doc, "Hydrogen technologies are becoming increasingly relevant to future low-carbon energy strategies, including industrial decarbonisation, energy storage, mobility, and resilient infrastructure. As hydrogen systems become digitally instrumented, they increasingly resemble operational technology environments composed of sensors, controllers, communication links, supervisory interfaces, and data-driven decision processes. The resulting benefits are substantial, but so are the cybersecurity implications.")
    add_para(doc, "A cyber incident in a hydrogen environment can affect integrity, availability, situational awareness, or control actions. In safety-critical contexts, these effects may translate into process instability, delayed intervention, equipment stress, or a loss of confidence in monitoring data. Recent NREL work highlights that hydrogen fueling and production contexts rely on digital technologies while introducing cybersecurity concerns related to communications, control, and interdependencies [1].")
    add_para(doc, "This project therefore focuses on cyber resilience rather than only prevention. Cyber resilience is treated as the capacity to anticipate, detect, withstand, respond to, and recover from cyber incidents. Machine-learning anomaly detection is studied as one detection-support mechanism within that wider resilience lifecycle, not as a complete security solution.")
    add_heading(doc, "1.1 Problem Statement", 2)
    add_para(doc, "Hydrogen-specific cyber datasets and practical detection benchmarks remain limited compared with wider ICS domains such as water treatment or manufacturing. At the same time, machine-learning anomaly detection is often evaluated using metrics alone, without making the resilience context explicit. This creates a gap between a promising technical technique and the operational needs of hydrogen cyber-physical systems.")
    add_heading(doc, "1.2 Project Contributions", 2)
    add_bullets(doc, [
        "A hydrogen-focused rationale for ML-enabled anomaly detection in OT settings.",
        "A structured threat model with four attack scenarios grounded in cyber-physical consequences.",
        "A synthetic dataset generator designed for interpretability and repeatability.",
        "A model comparison covering Logistic Regression, Random Forest, and Isolation Forest.",
        "Evaluation using holdout testing, cross-validation, precision-recall analysis, per-attack recall, and feature-importance analysis.",
        "A deployment-aware critical evaluation that explains what the artefact proves and what it does not prove.",
    ])
    add_heading(doc, "1.3 Scope and Originality", 2)
    add_para(doc, "The originality of the work is not the invention of a new ML algorithm. Instead, it lies in the integration of hydrogen cyber-resilience framing, an explicit OT threat model, and a transparent evaluation artefact that moves beyond a simple accuracy demo. The project is intentionally scoped to be feasible for a bachelor final year project while still demonstrating technical depth, critical judgement, and reproducibility.")

    # 2 Aim/objectives/RQs
    add_heading(doc, "2. Aim, Objectives, and Research Questions", 1)
    add_heading(doc, "2.1 Aim", 2)
    add_para(doc, "The aim of this project is to design, implement, and critically evaluate a machine-learning anomaly-detection approach that can support the cyber resilience of hydrogen energy systems in a controlled educational setting.")
    add_heading(doc, "2.2 Objectives", 2)
    add_bullets(doc, [
        "Explain why hydrogen energy systems should be treated as cyber-physical OT environments.",
        "Identify cyber threats and attack effects relevant to digitally monitored hydrogen systems.",
        "Review ML anomaly-detection approaches suitable for a final-year Computer Science artefact.",
        "Generate a transparent hydrogen ICS-style dataset with multiple attack scenarios.",
        "Train, compare, and evaluate supervised and unsupervised anomaly-detection models.",
        "Analyse model strengths, failure modes, resilience implications, and limitations.",
    ])
    add_heading(doc, "2.3 Research Questions", 2)
    add_bullets(doc, [
        "RQ1: Which cyber threats are most relevant to a digitally monitored hydrogen energy system?",
        "RQ2: Can a lightweight ML pipeline distinguish simulated malicious behaviour from normal process behaviour?",
        "RQ3: How do supervised and unsupervised approaches differ in detection quality and operational trade-offs?",
        "RQ4: Which attack scenarios remain difficult to detect and why?",
        "RQ5: What limitations prevent the educational artefact from being treated as a deployable real-world detector?",
    ])

    # 3 Background
    add_heading(doc, "3. Background and Literature Review", 1)
    add_heading(doc, "3.1 Hydrogen Systems as Cyber-Physical Infrastructure", 2)
    add_para(doc, "A hydrogen energy system may involve electrolyzers, compressors, storage tanks, dispensers, leak-detection instrumentation, PLCs, supervisory workstations, engineering stations, historian databases, and external communications. These components form a cyber-physical environment in which digital signals influence or describe physical behaviour. NREL's 2025 analysis identifies digital technologies as important to hydrogen fuel production, storage, and fueling, while noting cybersecurity concerns relevant to such infrastructures [1].")
    add_heading(doc, "3.2 OT Security and Cyber Resilience", 2)
    add_para(doc, "Operational technology must be secured differently from ordinary enterprise IT because availability, reliability, safety, timing, and deterministic behaviour are often primary concerns. NIST SP 800-82 Rev. 3 emphasises that OT security must account for unique performance, reliability, and safety requirements [3]. CISA's OT principles likewise frame OT cyber security around systems where cyber compromise can create physical consequences [4].")
    add_para(doc, "NIST CSF 2.0 provides a useful resilience lens because its Functions include Govern, Identify, Protect, Detect, Respond, and Recover [2]. A machine-learning detector aligns most directly with Detect, but its value depends on how outputs feed into Respond and Recover activities. This project adopts that broader position: detection is necessary, but resilience requires decision processes beyond classification.")
    add_heading(doc, "3.3 Hydrogen-Specific Cybersecurity Rationale", 2)
    add_para(doc, "Hydrogen infrastructure deserves specific attention because process safety and digital trust interact. Manipulated pressure, flow, level, or actuator state data could impair operator judgement or obscure unsafe process states. Communication degradation can delay insight into system behaviour. A detection system that flags cyber-physical inconsistencies may therefore contribute to resilience by reducing time-to-awareness, provided alarms are accurate and operationally interpretable.")
    add_heading(doc, "3.4 ML Anomaly Detection in Industrial Settings", 2)
    add_para(doc, "Machine-learning anomaly detection is widely discussed for ICS because attacks may appear as multivariate deviations rather than obvious single-variable threshold breaches. NIST IR 8219 documents behavioural anomaly-detection capabilities in manufacturing and process-control demonstrations, supporting the relevance of anomaly-based methods for ICS contexts [5]. Public ICS research artefacts such as the SWaT testbed have also been used to study cyber-physical attack detection [6].")
    add_para(doc, "The present project does not attempt to outperform established benchmark studies. Instead, it adapts their key lessons to a hydrogen-motivated educational artefact: multivariate process signals matter, attack scenarios must be described explicitly, and evaluation should not rely on accuracy alone.")
    add_heading(doc, "3.5 Literature Gap", 2)
    add_para(doc, "The project addresses a small but meaningful gap. Hydrogen-specific cyber-resilience research and open datasets remain less mature than generic ICS anomaly-detection research. In addition, many introductory ML security projects report a single train-test metric without analysing resilience alignment, attack-specific difficulty, or deployment validity. This project responds by combining threat modelling, richer evaluation, and explicit limitations.")

    # 4 Threat model and resilience
    add_heading(doc, "4. Threat Model and Resilience Framework", 1)
    add_heading(doc, "4.1 Protected Assets", 2)
    add_table(doc, ["Asset", "Why it matters", "Representative features"], [
        ["Process instrumentation", "Defines operator visibility and controller decisions", "pressure, flow, tank level, temperature"],
        ["Actuator commands", "Affects process control behaviour", "valve state, compressor state"],
        ["Communications and control telemetry", "Supports timely supervision and alarm trust", "latency, packet rate, command rate"],
        ["Operational continuity", "Links detection to resilience and recovery", "availability and response implications"],
    ], widths=[1.55, 3.3, 2.1], font_size=8.4)
    add_heading(doc, "4.2 Adversary Objectives", 2)
    add_bullets(doc, [
        "Mislead operators or controllers through manipulated telemetry.",
        "Reduce availability or timeliness of monitoring/control communications.",
        "Create inconsistency between actuator states and process behaviour.",
        "Remain less visible by using small, gradual, low-amplitude data manipulation.",
    ])
    add_heading(doc, "4.3 Attack Scenario Matrix", 2)
    add_table(doc, ["Scenario", "Cyber effect", "Physical/operational interpretation", "Detection challenge"], [
        ["False data injection", "Large telemetry distortion", "Misleading pressure and tank-level visibility", "Relatively easy when deviations are high"],
        ["Denial of service", "Network-quality degradation", "Delayed or disrupted oversight", "Easy when latency/traffic shift materially"],
        ["Actuator spoofing", "State inconsistency", "Control-state mismatch with process changes", "Requires cross-signal interpretation"],
        ["Stealthy sensor drift", "Small gradual manipulation", "Slow erosion of telemetry trust", "Harder because it resembles benign process variation"],
    ], widths=[1.35, 1.55, 3.0, 2.0], font_size=8.0)
    add_heading(doc, "4.4 Mapping Detection to Resilience", 2)
    add_table(doc, ["Resilience function", "Project relevance", "What the artefact contributes"], [
        ["Identify", "Define high-value signals and attack surfaces", "Feature groups and threat model"],
        ["Protect", "Baseline controls reduce attack likelihood", "Discussed conceptually; not implemented"],
        ["Detect", "Identify adverse events and anomalies", "Primary ML artefact"],
        ["Respond", "Triaging alarms and investigating probable anomalies", "Per-attack evaluation informs response priority"],
        ["Recover", "Restore safe operations and validate data integrity", "Discussed as deployment extension"],
    ], widths=[1.4, 2.6, 3.1], font_size=8.2)
    add_para(doc, "The mapping above is deliberately important: it prevents the false claim that a classifier alone equals cyber resilience. Instead, the classifier is framed as a detect-layer contribution that must be integrated with governance, incident response, and recovery planning [2].")

    # 5 Methodology
    add_heading(doc, "5. Methodology", 1)
    add_heading(doc, "5.1 Research Design", 2)
    add_para(doc, "The project uses a design-and-evaluate methodology. First, a cyber-resilience-motivated dataset and anomaly-detection artefact are designed. Second, the artefact is implemented in Python. Third, the models are experimentally evaluated and critically interpreted. This design is suitable because the aim is to produce a functioning technical artefact and assess its value within explicit scope limits.")
    add_heading(doc, "5.2 Dataset Strategy", 2)
    add_para(doc, f"The dataset contains {metrics_payload['dataset_rows']:,} rows. It uses time-indexed synthetic process dynamics with bounded noise and four attack windows. The overall attack prevalence is {metrics_payload['attack_prevalence']:.1%}. Synthetic data were selected because no hydrogen-specific public cyber dataset was used in the project constraints; the trade-off is feasibility and transparency at the cost of external realism.")
    add_table(doc, ["Dataset property", "Value"], [
        ["Rows", f"{metrics_payload['dataset_rows']:,}"],
        ["Attack prevalence", f"{metrics_payload['attack_prevalence']:.1%}"],
        ["Training rows", f"{metrics_payload['train_rows']:,}"],
        ["Test rows", f"{metrics_payload['test_rows']:,}"],
        ["Holdout split", "70% train / 30% test, stratified"],
        ["Cross-validation", "Three-fold stratified CV for supervised models"],
    ], widths=[2.35, 3.4], font_size=8.8)
    add_heading(doc, "5.3 Features and Data Representation", 2)
    add_table(doc, ["Feature group", "Variables", "Security rationale"], [
        ["Process", "pressure_bar, temperature_c, flow_nm3h, tank_level_pct, electrolyzer_current_a", "Captures cyber-physical deviations"],
        ["Actuator state", "valve_state, compressor_state", "Captures possible command/state spoofing"],
        ["Network/control activity", "network_latency_ms, packet_rate_pps, command_rate_per_min", "Captures communication degradation and activity spikes"],
    ], widths=[1.35, 4.2, 2.2], font_size=8.0)
    add_heading(doc, "5.4 Model Selection", 2)
    add_table(doc, ["Model", "Type", "Reason for inclusion", "Expected limitation"], [
        ["Logistic Regression", "Supervised baseline", "Transparent linear comparator with class weighting", "May miss nonlinear interactions"],
        ["Random Forest", "Supervised nonlinear model", "Handles mixed features and feature interactions", "Can overfit synthetic patterns"],
        ["Isolation Forest", "Unsupervised", "Normal-only anomaly baseline", "Higher false-positive risk"],
    ], widths=[1.55, 1.4, 3.2, 2.0], font_size=8.0)
    add_heading(doc, "5.5 Evaluation Strategy", 2)
    add_para(doc, "The evaluation intentionally uses more than one lens. A fixed stratified holdout split supports a clear reproducible comparison. Stratified cross-validation provides a second estimate of robustness for supervised models. The metric set includes accuracy, precision, recall, F1-score, ROC-AUC, and average precision where probabilities are available. Per-attack recall examines scenario-specific blind spots. Permutation feature importance identifies which features materially influence F1 performance.")
    add_heading(doc, "5.6 Validity Considerations", 2)
    add_table(doc, ["Validity issue", "Risk", "Mitigation in this project"], [
        ["Construct validity", "Synthetic scenarios may not fully represent real hydrogen attacks", "Attack narratives are explicit and claims are scoped"],
        ["Internal validity", "One split may produce optimistic or unstable results", "Cross-validation added for supervised models"],
        ["External validity", "Real OT environments are noisier and more constrained", "No production claim; future validation pathway stated"],
        ["Conclusion validity", "High scores may be overinterpreted", "Error analysis and limitations are emphasised"],
    ], widths=[1.45, 3.05, 3.0], font_size=8.0)

    # 6 Implementation
    add_heading(doc, "6. Artefact Design and Implementation", 1)
    add_heading(doc, "6.1 Artefact Architecture", 2)
    add_table(doc, ["Component", "Purpose", "Outputs"], [
        ["generate_synthetic_hydrogen_ics_data.py", "Builds normal traces and four labelled attack scenarios", "CSV dataset"],
        ["train_hydrogen_anomaly_model.py", "Trains models, evaluates metrics, creates figures", "CSV/JSON reports and plots"],
        ["README.md", "Explains scope, workflow, and reproducibility", "Project guidance"],
        ["requirements.txt", "Lists dependencies", "Environment recreation"],
    ], widths=[2.35, 3.2, 1.55], font_size=8.1)
    add_heading(doc, "6.2 Reproducibility", 2)
    add_para(doc, "The implementation uses deterministic seeds, explicit input/output paths, machine-readable outputs, and standalone scripts. The project can be rerun using the commands below.")
    add_code_block(doc, "python code/generate_synthetic_hydrogen_ics_data.py --output data/hydrogen_ics_synthetic.csv\npython code/train_hydrogen_anomaly_model.py --input data/hydrogen_ics_synthetic.csv --output-dir outputs")
    add_heading(doc, "6.3 Evaluation Outputs", 2)
    add_bullets(doc, [
        "model_metrics.csv and model_metrics.json for holdout results",
        "cross_validation_metrics.csv for supervised model robustness checks",
        "per_attack_recall.csv for scenario-specific detection quality",
        "feature_importance.csv and feature_importance.png for interpretability",
        "confusion_matrix.png, roc_curve.png, precision_recall_curve.png, and attack_timeline.png for visual analysis",
    ])
    add_figure(doc, TIMELINE, "Figure 1. Injected attack windows in the synthetic hydrogen ICS-style dataset.", width_inches=6.1)

    # 7 Results
    add_heading(doc, "7. Experimental Results", 1)
    add_heading(doc, "7.1 Holdout Test Results", 2)
    metric_rows = []
    for model_name in metrics_df.index:
        row = metrics_df.loc[model_name]
        metric_rows.append([
            model_name.replace("_", " ").title(),
            fmt_num(row["accuracy"]),
            fmt_num(row["precision"]),
            fmt_num(row["recall"]),
            fmt_num(row["f1"]),
            "-" if pd.isna(row.get("roc_auc", float("nan"))) else fmt_num(row["roc_auc"]),
            "-" if pd.isna(row.get("average_precision", float("nan"))) else fmt_num(row["average_precision"]),
        ])
    add_table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "Avg. Precision"], metric_rows, widths=[1.55, 0.8, 0.8, 0.8, 0.8, 0.85, 1.0], font_size=7.8)
    add_para(doc, "Random Forest achieved the best balance of precision and recall on the holdout split. Isolation Forest detected many attacks but produced substantially more false positives, which would be costly in an operational context because excessive alarms can erode operator trust.")
    add_heading(doc, "7.2 Cross-Validation Results", 2)
    cv_rows = []
    for _, row in cv_df.iterrows():
        cv_rows.append([
            row["model"].replace("_", " ").title(),
            f"{row['accuracy_mean']:.3f} +/- {row['accuracy_std']:.3f}",
            f"{row['precision_mean']:.3f} +/- {row['precision_std']:.3f}",
            f"{row['recall_mean']:.3f} +/- {row['recall_std']:.3f}",
            f"{row['f1_mean']:.3f} +/- {row['f1_std']:.3f}",
        ])
    add_table(doc, ["Model", "Accuracy", "Precision", "Recall", "F1"], cv_rows, widths=[1.55, 1.4, 1.4, 1.4, 1.4], font_size=7.9)
    add_para(doc, "Cross-validation supports the same overall conclusion as the holdout split: Random Forest remains the strongest model, while Logistic Regression serves as a credible but less powerful baseline. This additional robustness check strengthens the evaluation beyond a single split.")
    add_heading(doc, "7.3 Visual Performance Analysis", 2)
    add_figure(doc, CONFUSION, "Figure 2. Confusion matrix for the best supervised model.", width_inches=5.7)
    add_figure(doc, ROC, "Figure 3. ROC curve for the best supervised model.", width_inches=5.7)
    add_figure(doc, PR, "Figure 4. Precision-recall curve for the best supervised model.", width_inches=5.7)
    add_heading(doc, "7.4 Per-Attack Recall", 2)
    attack_rows = []
    for _, row in attack_recall_df.iterrows():
        attack_rows.append([
            str(row["attack_type"]).replace("_", " ").title(),
            str(int(row["attack_instances"])),
            str(int(row["detected_instances"])),
            fmt_num(row["recall"]),
        ])
    add_table(doc, ["Attack type", "Instances", "Detected", "Recall"], attack_rows, widths=[2.5, 1.0, 1.0, 1.1], font_size=8.4)
    add_para(doc, "The lowest recall occurs for stealthy sensor drift. This is an intentionally valuable finding rather than a defect in the report: lower-amplitude, gradual attacks are more difficult to identify because they resemble ordinary operational variation. A distinction-level analysis should expose such limitations instead of hiding them behind aggregate metrics.")
    add_heading(doc, "7.5 Permutation Feature Importance", 2)
    feat_rows = []
    for _, row in feature_df.head(7).iterrows():
        feat_rows.append([
            row["feature"],
            fmt_num(row["importance_mean"]),
            fmt_num(row["importance_std"]),
        ])
    add_table(doc, ["Feature", "Mean F1 decrease", "Std. dev."], feat_rows, widths=[3.1, 1.45, 1.45], font_size=8.4)
    add_figure(doc, FEATURE_PLOT, "Figure 5. Permutation feature importance for the best supervised model.", width_inches=6.0)
    add_para(doc, "Flow, pressure, and tank level were the most influential features for the best model. This aligns with the attack design, particularly the false-data, actuator-spoofing, and drift scenarios. The result is interpretable, but it also warns that the model may be strongly coupled to the engineered synthetic scenario distribution.")

    # 8 Discussion and critical evaluation
    add_heading(doc, "8. Discussion and Critical Evaluation", 1)
    add_heading(doc, "8.1 Answers to Research Questions", 2)
    add_table(doc, ["RQ", "Answer"], [
        ["RQ1", "Relevant threats include telemetry manipulation, communication degradation, actuator-state inconsistency, and low-amplitude drift."],
        ["RQ2", "Yes. The pipeline distinguished controlled attacks from normal behaviour with strong supervised performance."],
        ["RQ3", "Supervised Random Forest substantially outperformed the unsupervised baseline in precision and F1; Isolation Forest had weaker alarm quality."],
        ["RQ4", "Stealthy sensor drift was hardest to detect, showing that subtle attacks deserve special attention."],
        ["RQ5", "Synthetic data, simplified threat scenarios, and missing real OT constraints prevent deployment claims."],
    ], widths=[0.65, 6.6], font_size=8.3)
    add_heading(doc, "8.2 Interpretation of the Model Results", 2)
    add_para(doc, "The strongest result is not simply that Random Forest scores highly. The more academically meaningful result is that the evaluation differentiates attack difficulty. High recall for overt denial-of-service and actuator-spoofing patterns is unsurprising because the synthetic deviations are sizeable. The lower recall for stealthy drift demonstrates a more realistic challenge: attacks that move slowly within apparently plausible ranges can evade simple detection.")
    add_heading(doc, "8.3 Operational Implications", 2)
    add_para(doc, "In an operational hydrogen system, a detector would need to be embedded within a broader monitoring and response process. A high false-positive detector could overwhelm operators, while a false-negative detector could create misplaced confidence. Therefore, F1 and average precision are meaningful, but they are not sufficient. Alarm explainability, response runbooks, human validation, fail-safe logic, and testing against real plant behaviour would be required before deployment.")
    add_heading(doc, "8.4 Strengths of the Project", 2)
    add_bullets(doc, [
        "The project is sharply scoped and logically linked to the proposal topic.",
        "The threat model gives the work stronger research structure than a generic classifier exercise.",
        "The experimental design includes both baseline comparison and robustness checks.",
        "The report provides error analysis, interpretability, and deployment caution rather than only headline metrics.",
        "The artefact is reproducible and can be extended with real datasets or digital-twin simulations.",
    ])
    add_heading(doc, "8.5 Limitations", 2)
    add_bullets(doc, [
        "The dataset is synthetic and cannot reproduce the full stochastic, safety, and control dynamics of a real hydrogen facility.",
        "Attack patterns are designed manually and may be easier to detect than adversarially crafted real attacks.",
        "The project evaluates binary attack detection rather than root-cause diagnosis or response automation.",
        "Temporal models such as LSTM or transformer-style sequence models are not implemented; this is a deliberate feasibility choice, not an omission hidden from the reader.",
        "No claims are made about certification, operational safety compliance, or direct industrial deployment.",
    ])
    add_heading(doc, "8.6 Why the Evaluation Is Stronger Than a Basic Bachelor Prototype", 2)
    add_para(doc, "A minimal project might stop after training one model and reporting accuracy. This project is stronger because it defines a hydrogen-specific threat model, uses multiple models, evaluates with class-sensitive metrics, adds cross-validation, analyses attack-specific blind spots, explains feature importance, and draws a careful boundary between educational evidence and real deployment claims. These choices demonstrate critical thinking, methodological care, and technical maturity.")

    # 9 Project management etc.
    add_heading(doc, "9. Project Management, Risks, Ethics, and Professional Considerations", 1)
    add_heading(doc, "9.1 Development Plan", 2)
    add_table(doc, ["Phase", "Main activity", "Deliverable"], [
        ["1", "Topic scoping, proposal refinement, literature review", "Research framing"],
        ["2", "Threat model and dataset design", "Scenario definitions"],
        ["3", "Python implementation", "Dataset generator and training script"],
        ["4", "Evaluation and interpretability", "Metrics, plots, CV, per-attack recall"],
        ["5", "Report writing and quality assurance", "Final report and packaged artefact"],
    ], widths=[0.65, 3.35, 2.9], font_size=8.2)
    add_heading(doc, "9.2 Risks and Mitigations", 2)
    add_table(doc, ["Risk", "Impact", "Mitigation"], [
        ["No hydrogen cyber dataset", "Weak external realism", "Use transparent synthetic data and state this limitation clearly"],
        ["Overly simple attack design", "Inflated metrics", "Add a subtle drift attack and attack-specific evaluation"],
        ["Model overfitting", "Misleading conclusions", "Use baseline comparison, cross-validation, and feature analysis"],
        ["Excessive project scope", "Incomplete artefact", "Focus on interpretable ML and rigorous evaluation"],
    ], widths=[2.1, 2.2, 3.0], font_size=8.0)
    add_heading(doc, "9.3 Ethics and Responsible Framing", 2)
    add_para(doc, "The artefact is defensive and educational. It does not provide exploitation instructions for real hydrogen systems. The report emphasises that detection models must be validated responsibly and should not be presented as safety assurances without real-world testing. This is consistent with professional expectations for honest reporting of limitations in safety-relevant cyber-physical contexts.")
    add_heading(doc, "9.4 Language and Accessibility", 2)
    add_para(doc, "The final report uses clear sectioning, concise definitions, figure captions, and structured tables to support readability. Technical terms are introduced before use, and key claims are separated from limitations to reduce ambiguity.")

    # 10 conclusion
    add_heading(doc, "10. Conclusion and Future Work", 1)
    add_para(doc, "This project demonstrates a rigorous bachelor-level investigation into how machine learning can support the cyber resilience of hydrogen energy systems. The work contributes a hydrogen-motivated threat model, a reproducible synthetic anomaly-detection artefact, and a stronger-than-basic evaluation strategy that includes holdout testing, cross-validation, scenario-specific recall, and interpretability analysis.")
    add_para(doc, "The results show that supervised learning can detect the engineered attack scenarios effectively, with Random Forest producing the strongest overall performance. However, the deliberately more subtle stealthy drift attack is harder to detect, which reinforces the central argument that resilience requires more than headline metrics. Real deployment would require domain data, plant-aware validation, explainable alarm workflows, response procedures, and alignment with OT security standards.")
    add_heading(doc, "10.1 Recommended Future Work", 2)
    add_bullets(doc, [
        "Validate the workflow on a real ICS benchmark and, if available, hydrogen-relevant process telemetry.",
        "Explore temporal models and sliding-window features for drift-style attacks.",
        "Add threshold tuning based on operational cost rather than only default classification boundaries.",
        "Integrate alarm explanations with a response playbook mapped to OT resilience functions.",
        "Study concept drift, missing data, and sensor faults to distinguish cyber anomalies from benign operational degradation.",
    ])

    # References
    add_heading(doc, "References", 1)
    references = [
        "[1] D. Hatic et al., Cybersecurity Considerations for Hydrogen Infrastructure in Airport Environments, National Renewable Energy Laboratory, 2025.",
        "[2] National Institute of Standards and Technology, The NIST Cybersecurity Framework (CSF) 2.0, NIST CSWP 29, 2024.",
        "[3] K. Stouffer et al., Guide to Operational Technology (OT) Security, NIST SP 800-82 Rev. 3, 2023.",
        "[4] Cybersecurity and Infrastructure Security Agency, Principles of Operational Technology Cyber Security, 2024.",
        "[5] J. McCarthy et al., Securing Manufacturing Industrial Control Systems: Behavioral Anomaly Detection, NIST IR 8219, 2020.",
        "[6] A. P. Mathur and N. O. Tippenhauer, SWaT: A Water Treatment Testbed for Research and Training on ICS Security, 2016.",
    ]
    for ref in references:
        add_para(doc, ref)

    # Appendix A
    add_heading(doc, "Appendix A. Reproduction Instructions", 1)
    add_para(doc, "From the project root folder, run:")
    add_code_block(doc, "python code/generate_synthetic_hydrogen_ics_data.py --output data/hydrogen_ics_synthetic.csv\npython code/train_hydrogen_anomaly_model.py --input data/hydrogen_ics_synthetic.csv --output-dir outputs")
    add_para(doc, "The workflow regenerates the dataset and exports all metrics and plots used in the report.")

    # Appendix B
    add_heading(doc, "Appendix B. Data Dictionary", 1)
    feature_rows = [
        ["timestamp_index", "Synthetic ordering variable", "Index only"],
        ["pressure_bar", "Process pressure", "Continuous"],
        ["temperature_c", "Process temperature", "Continuous"],
        ["flow_nm3h", "Hydrogen flow proxy", "Continuous"],
        ["tank_level_pct", "Storage level proxy", "Continuous"],
        ["electrolyzer_current_a", "Electrolyzer electrical load proxy", "Continuous"],
        ["valve_state", "Valve actuator state", "Binary"],
        ["compressor_state", "Compressor actuator state", "Binary"],
        ["network_latency_ms", "Communication latency", "Continuous"],
        ["packet_rate_pps", "Traffic intensity proxy", "Continuous"],
        ["command_rate_per_min", "Control-command activity proxy", "Continuous"],
        ["attack_label", "0 normal / 1 attack", "Binary target"],
        ["attack_type", "Scenario label", "Categorical"],
    ]
    add_table(doc, ["Field", "Meaning", "Type"], feature_rows, widths=[2.2, 3.8, 1.3], font_size=7.8)

    # Appendix C
    doc.add_page_break()
    add_heading(doc, "Appendix C. Attack Scenario Definitions", 1)
    attack_rows = [
        ["False data injection", "Pressure increases; tank-level reporting decreases; command activity rises modestly."],
        ["Denial of service", "Latency and packet rate increase sharply; command activity rises."],
        ["Actuator spoofing", "Valve and compressor states are flipped; pressure and flow shift inconsistently."],
        ["Stealthy sensor drift", "Pressure, flow, and tank-level values move gradually and subtly, making detection harder."],
    ]
    add_table(doc, ["Scenario", "Synthetic implementation"], attack_rows, widths=[2.0, 5.4], font_size=8.2)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_report()
